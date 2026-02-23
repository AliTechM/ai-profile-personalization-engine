"""
DOCX export utilities for the Resume schema.

Visually matches HTML template:
- Section separation
- Left title / right bold dates
- No italics anywhere
"""

from __future__ import annotations

from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.normalization import normalize_resume_for_template
from schemas.resume import Resume


# =========================
# Layout constants
# =========================
SECTION_SPACE_BEFORE = Pt(18)
SECTION_SPACE_AFTER = Pt(6)
ENTRY_SPACE_BEFORE = Pt(10)
ENTRY_SPACE_AFTER = Pt(2)
BULLET_INDENT = Pt(18)
RIGHT_MARGIN_POSITION = Pt(450)  # right-aligned tab stop


# =========================
# Helpers
# =========================

def _add_section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SECTION_SPACE_BEFORE
    p.paragraph_format.space_after = SECTION_SPACE_AFTER

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)

    # bottom border (section separator)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")

    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_left_right_entry(
    doc: Document,
    left_text: str,
    right_text: str | None = None,
) -> None:
    """
    Left-aligned title, right-aligned bold date on SAME LINE
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = ENTRY_SPACE_BEFORE
    p.paragraph_format.space_after = ENTRY_SPACE_AFTER

    # Right-aligned tab stop
    p.paragraph_format.tab_stops.add_tab_stop(
        RIGHT_MARGIN_POSITION,
        WD_TAB_ALIGNMENT.RIGHT,
    )

    run_left = p.add_run(left_text)
    run_left.bold = True

    if right_text:
        p.add_run("\t")
        run_right = p.add_run(right_text)
        run_right.bold = True


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = BULLET_INDENT
    p.paragraph_format.space_after = Pt(2)


# =========================
# Renderer
# =========================

def render_resume_docx(resume: Resume) -> bytes:
    context = normalize_resume_for_template(resume)
    doc = Document()

    # =========================
    # Header
    # =========================
    personal = context.get("personal_info", {})

    if personal.get("full_name"):
        p = doc.add_paragraph()
        run = p.add_run(personal["full_name"])
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = 1

    contact = " | ".join(
        c for c in [
            personal.get("email_address"),
            personal.get("phone_number"),
            personal.get("linkedin"),
        ] if c
    )
    if contact:
        p = doc.add_paragraph(contact)
        p.alignment = 1

    if personal.get("personal_website"):
        p = doc.add_paragraph(personal["personal_website"])
        p.alignment = 1

    doc.add_paragraph().paragraph_format.space_after = Pt(9)

    # =========================
    # Professional Summary
    # =========================
    if context.get("summary"):
        _add_section_title(doc, "Professional Summary")
        doc.add_paragraph(context["summary"])

    # =========================
    # Education
    # =========================
    educations = context.get("educations", [])
    if educations:
        _add_section_title(doc, "Education")
        for edu in educations:
            title = f"{edu.get('degree')} in {edu.get('major')}"
            dates = f"{edu.get('start_date')} - {edu.get('end_date')}".strip(" -")

            _add_left_right_entry(doc, title, dates)

            university = ", ".join(
                p for p in [
                    edu.get("university_name"),
                    edu.get("city"),
                    edu.get("country"),
                ] if p
            )
            if university:
                doc.add_paragraph(university)

    # =========================
    # Professional Experience
    # =========================
    professional_exps = context.get("professional_experiences", [])
    if professional_exps:
        _add_section_title(doc, "Professional Experience")
        for exp in professional_exps:
            dates = f"{exp.get('start_date')} - {exp.get('end_date')}".strip(" -")
            _add_left_right_entry(
                doc,
                exp.get("role_title", ""),
                dates,
            )

            if exp.get("company_name"):
                doc.add_paragraph(exp["company_name"])

            for bullet in exp.get("description", []):
                _add_bullet(doc, bullet)

    # =========================
    # Volunteer Experience
    # =========================
    volunteer_exps = context.get("volunteer_experiences", [])
    if volunteer_exps:
        _add_section_title(doc, "Volunteer Experience")
        for exp in volunteer_exps:
            dates = f"{exp.get('start_date')} - {exp.get('end_date')}".strip(" -")
            _add_left_right_entry(
                doc,
                exp.get("role_title", ""),
                dates,
            )

            if exp.get("company_name"):
                doc.add_paragraph(exp["company_name"])

            for bullet in exp.get("description", []):
                _add_bullet(doc, bullet)

    # =========================
    # Projects
    # =========================
    projects = context.get("projects", [])
    if projects:
        _add_section_title(doc, "Projects")
        for project in projects:
            _add_left_right_entry(doc, project.get("project_name", ""))

            for bullet in project.get("description", []):
                _add_bullet(doc, bullet)

            if project.get("project_link"):
                doc.add_paragraph(f"Project link: {project['project_link']}")

    # =========================
    # Certifications
    # =========================
    certifications = context.get("certifications", [])
    if certifications:
        _add_section_title(doc, "Certifications")
        for cert in certifications:
            doc.add_paragraph(
                " | ".join(
                    p for p in [
                        cert.get("certification_name"),
                        cert.get("issuing_organization"),
                        cert.get("issue_date"),
                    ] if p
                )
            )

    # =========================
    # Skills
    # =========================
    skills = context.get("skills_grouped", {})
    if skills:
        _add_section_title(doc, "Core Competencies")
        if skills.get("technical"):
            # Heading line
            p = doc.add_paragraph()
            run = p.add_run("Technical Skills")
            run.bold = True

            # Content line
            doc.add_paragraph(", ".join(skills["technical"]))
        if skills.get("soft"):
            # Heading line
            p = doc.add_paragraph()
            run = p.add_run("Soft Skills")
            run.bold = True

            # Content line
            doc.add_paragraph(", ".join(skills["soft"]))

    # =========================
    # Languages
    # =========================
    languages = context.get("languages", [])
    if languages:
        _add_section_title(doc, "Languages")
        for lang in languages:
            doc.add_paragraph(
                f"{lang.get('language')}: {lang.get('proficiency_level')}"
            )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
